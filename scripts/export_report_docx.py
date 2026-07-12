"""Export B9AI103 CA01 report markdown to Word (.docx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
MAIN_MD = REPORT_DIR / "B9AI103_CA01_Report.md"
APPENDICES = [
    REPORT_DIR / "appendices" / "Appendix_A_Normalisation.md",
    REPORT_DIR / "appendices" / "Appendix_B_Database_Schema.md",
]
OUTPUT = REPORT_DIR / "B9AI103_CA01_Report.docx"


def set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_runs_with_inline_format(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            cell_obj = table.rows[r_idx].cells[c_idx]
            cell_obj.text = ""
            p = cell_obj.paragraphs[0]
            add_runs_with_inline_format(p, cell)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_image_if_exists(doc: Document, alt: str, src: str, base_dir: Path) -> None:
    candidates = [
        base_dir / src,
        REPORT_DIR / "appendices" / src,
        REPORT_DIR / src,
    ]
    for path in candidates:
        if path.exists():
            doc.add_paragraph(alt, style="Caption")
            doc.add_picture(str(path), width=Inches(5.5))
            return
    p = doc.add_paragraph()
    p.add_run(f"[Image: {alt} — file not found: {src}]").italic = True


def render_markdown(doc: Document, md_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    list_buffer: list[tuple[str, str]] = []

    def flush_list() -> None:
        nonlocal list_buffer
        for kind, item in list_buffer:
            style = "List Number" if kind == "num" else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_runs_with_inline_format(p, item)
        list_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_list()
            in_code = True
            code_lines = []
            i += 1
            continue

        if not line.strip():
            flush_list()
            i += 1
            continue

        if line.strip() == "---":
            flush_list()
            doc.add_paragraph()
            i += 1
            continue

        img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img:
            flush_list()
            add_image_if_exists(doc, img.group(1), img.group(2), md_path.parent)
            i += 1
            continue

        if line.lstrip().startswith("|") and "|" in line:
            flush_list()
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                if not is_table_separator(lines[i]):
                    table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        num = re.match(r"^(\d+)\.\s+(.*)$", line)
        if num:
            list_buffer.append(("num", num.group(2)))
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", line)
        if bullet:
            list_buffer.append(("bullet", bullet.group(1)))
            i += 1
            continue

        flush_list()

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            doc.add_heading(heading.group(2), level=level)
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_inline_format(p, line)
        i += 1

    flush_list()
    if in_code and code_lines:
        add_code_block(doc, code_lines)


def count_main_body_words(md_path: Path) -> int:
    text = md_path.read_text(encoding="utf-8")
    start = text.find("## 1. Introduction")
    end = text.find("## Appendices")
    if start == -1 or end == -1:
        body = text
    else:
        body = text[start:end]
    body = re.sub(r"```.*?```", " ", body, flags=re.S)
    body = re.sub(r"\|.*\|", " ", body)
    body = re.sub(r"[#*_`\[\]()>|-]", " ", body)
    words = re.findall(r"[A-Za-z0-9]+(?:'[A-Za-z]+)?", body)
    return len(words)


def main() -> None:
    doc = Document()
    set_document_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FridgeWise AI: A Hybrid Recipe Recommender System for Food Waste Reduction")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("B9AI103 Recommender Systems — CA ONE Group Report").italic = True

    doc.add_paragraph()
    word_count = count_main_body_words(MAIN_MD)
    note = doc.add_paragraph()
    note.add_run(f"Estimated main-body word count (sections 1–15, excluding appendices): {word_count}").italic = True
    doc.add_page_break()

    render_markdown(doc, MAIN_MD)

    for appendix in APPENDICES:
        if appendix.exists():
            doc.add_page_break()
            render_markdown(doc, appendix)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Main-body word count estimate: {word_count} (limit 3,000 excluding appendices)")


if __name__ == "__main__":
    main()
