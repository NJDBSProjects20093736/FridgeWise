"""Export ThriftyChef web app user guide to Word (.docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deploy" / "ThriftyChef_Web_App_User_Guide.docx"
OUTPUT_V5 = ROOT / "deploy" / "ThriftyChef_Web_App_User_Guide_v5.docx"
MARKDOWN = ROOT / "deploy" / "ThriftyChef_Web_App_User_Guide.md"


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x34, 0xA0, 0xA4)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    width = max(len(r) for r in rows) if rows else len(headers)
    rows = [r + [""] * (width - len(r)) for r in rows]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def build() -> None:
    doc = Document()
    set_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ThriftyChef — Web App User Guide")
    r.bold = True
    r.font.size = Pt(20)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Product prototype — Flutter web + FastAPI (v5 — camera barcode scan added)").italic = True
    doc.add_paragraph()
    doc.add_paragraph(
        "ThriftyChef is a smart fridge and recipe assistant that helps reduce food waste by "
        "tracking ingredients, expiry dates, dietary requirements, and delivering personalised "
        "AI recipe recommendations."
    )
    doc.add_paragraph(
        "Design: teal primary (#34A0A4), ThriftyChef wordmark with chef-hat on “Chef”, "
        "food photography on recipe cards, responsive layout. Screen set: deploy/ThriftyChef_Screen_Set.png"
    )

    doc.add_heading("1. Start the servers", level=1)
    doc.add_heading("Terminal 1 — API backend", level=2)
    code_block(
        doc,
        'cd "D:\\DBS - Sem 2\\RC\\Fridge-Wise"\n'
        ".\\.venv\\Scripts\\python.exe scripts\\run_api.py",
    )
    doc.add_paragraph("Wait for: Uvicorn running on http://0.0.0.0:8000")
    doc.add_paragraph("Health check: http://127.0.0.1:8000/health")

    doc.add_heading("Terminal 2 — Web app (release build — recommended)", level=2)
    doc.add_paragraph(
        "Use a release build served as static files. This avoids blank-screen issues with flutter run in debug mode."
    )
    code_block(
        doc,
        'cd "D:\\DBS - Sem 2\\RC\\Fridge-Wise\\app"\n'
        "flutter pub get\n"
        "flutter build web --release --dart-define=API_BASE_URL=http://127.0.0.1:8000\n"
        "cd build\\web\n"
        "..\\..\\..\\.venv\\Scripts\\python.exe -m http.server 8080 --bind 127.0.0.1",
    )
    add_table(
        doc,
        ["Service", "URL"],
        [
            ["Web app", "http://127.0.0.1:8080"],
            ["API docs", "http://127.0.0.1:8000/docs"],
            ["API health", "http://127.0.0.1:8000/health"],
        ],
    )
    doc.add_paragraph("Windows tip: use 127.0.0.1 instead of localhost for the web app on port 8080.")
    doc.add_paragraph(
        "After code changes: rebuild with flutter build web --release, then restart the Python server."
    )

    doc.add_heading("Docker (optional — Linux server)", level=2)
    doc.add_paragraph("See docker/README.md or deploy/SERVER_DEPLOY.md for the full Docker stack.")

    doc.add_heading("2. Loading & first launch", level=1)
    add_numbered(
        doc,
        [
            "Boot splash — “Loading ThriftyChef…” with teal spinner",
            "API health check — GET /health before main UI",
            "Onboarding on first visit; returning users go to Recipes tab",
        ],
    )

    doc.add_heading("3. App shell & navigation", level=1)
    add_table(
        doc,
        ["Element", "Description"],
        [
            ["Header", "ThriftyChef wordmark — teal “Thrifty”, rose “Chef” with chef-hat icon"],
            ["API status", "Green dot = connected · Red = offline/fallback"],
            ["Theme toggle", "Sun/moon icon — light and dark mode"],
            ["Profile icon", "Edit diet/allergies anytime"],
            ["Recipes", "AI hybrid recommendations (default tab)"],
            ["Fridge", "Inventory, expiry tracking, add-item form"],
            ["Scan", "Barcode lookup — camera scan or manual entry, fridge or rescue-basket demo"],
            ["Substitute", "Unfamiliar ingredient similarity (cold-start)"],
        ],
    )
    doc.add_paragraph(
        "Mobile/tablet: bottom navigation (4 tabs). Desktop: left navigation rail + centred content (~1140px)."
    )

    doc.add_heading("4. Light & dark mode", level=1)
    add_bullets(
        doc,
        [
            "Light: pale frost background, white cards, teal #34A0A4 accents",
            "Dark: deep navy background, teal glow on cards, rose “Chef” in logo",
            "Toggle via sun/moon icon in top app bar; preference saved locally",
        ],
    )

    doc.add_heading("5. First-time setup (Onboarding)", level=1)
    doc.add_paragraph(
        "Centred profile card (max ~720px) with logo and subtitle: "
        "“Personalise your recommendations before we search your fridge.”"
    )
    doc.add_paragraph("Sections A–E in cards:")
    add_bullets(
        doc,
        [
            "A. Dietary requirement: radio cards — none, vegetarian, vegan, halal (hard safety filter)",
            "B. Allergies: chips — milk, eggs, peanuts, gluten, soy, fish (hard safety filter)",
            "C. Nutrition preferences: low sugar, low fat, gluten free, high protein",
            "D. Cuisine preferences: Italian, Asian, Indian, Mexican, Mediterranean, Sri Lankan, Any",
            "E. Openness to new cuisines: slider 0.0–1.0",
        ],
    )
    doc.add_paragraph(
        "Selected chips turn teal. Button: “Save profile and continue” with loading state. "
        "Profile saved via PUT /users/5060/profile and locally as fallback."
    )

    doc.add_heading("6. Fridge inventory", level=1)
    doc.add_paragraph(
        "Desktop: two-column — item list left, Add ingredient form right. "
        "Mobile: scrollable list with add form below."
    )
    doc.add_paragraph("Summary cards: Total items, Expiring soon, Barcode products.")
    doc.add_paragraph("Filter chips: All, Expiring soon, Barcode items, Safe ingredients.")
    add_table(
        doc,
        ["Field", "Description"],
        [
            ["Ingredient name", "e.g. tomato, eggs, parsley"],
            ["Quantity / unit", "Optional, e.g. 2 pcs"],
            ["Days to expiry", "Colour-coded slider 1–30 days"],
            ["Barcode", "Camera scan or manual entry (e.g. 6111246721261)"],
        ],
    )
    doc.add_paragraph("Item cards show circular ingredient photos, expiry progress bars, and urgency colours:")
    add_bullets(
        doc,
        [
            "Red: 0–2 days to expiry",
            "Amber: 3–5 days",
            "Green: 6+ days",
        ],
    )
    doc.add_paragraph(
        "Barcode nutrition panel: product name, brand, Nutri-Score, metric tiles, allergen chips, "
        "“Add product to fridge”. Edit/delete with confirmation. Changes refresh AI recommendations."
    )

    doc.add_heading("7. AI Recommendations", level=1)
    doc.add_paragraph(
        "Main demo screen: hero card, context/model badges, mood chips, expiry/context toggles, search card."
    )
    add_bullets(
        doc,
        [
            "Mood chips: Comfort, Healthy, Quick, Adventurous, Celebration",
            "Use expiry priority: boosts recipes using items expiring soon",
            "Use context boost: season/weekday re-ranking",
            "Search bar: filter AI list or search full catalogue",
            "Tune menu: switch Hybrid / Content / Collaborative / Popularity",
        ],
    )
    doc.add_paragraph("Each recipe card shows:")
    add_bullets(
        doc,
        [
            "Food photo at top (Unsplash; placeholder if offline)",
            "Circular match % badge and AI score",
            "Safe shield badge (passes diet/allergy filters)",
            "Tags: expiring (orange), high match, quick, nutrition, missing count",
            "Prep time and short reason line",
        ],
    )
    doc.add_paragraph(
        "Responsive wrap layout — cards reflow on narrow screens. Featured carousel on wider screens."
    )

    doc.add_heading("8. Recipe detail", level=1)
    doc.add_paragraph("Tap any recipe card:")
    add_numbered(
        doc,
        [
            "Hero food image at top",
            "Summary card — match %, AI score, prep time, safety badge",
            "Why recommended — explainable AI bullets",
            "Ingredients and missing ingredient chips",
            "Method / cooking steps (numbered timeline)",
            "Nutrition notes, allergy safety, possible substitutions",
        ],
    )
    doc.add_paragraph("Desktop: two-column. Mobile: stacked with back button.")

    doc.add_heading("9. Scan (Barcode tab)", level=1)
    doc.add_paragraph(
        "Scan food — check recipes before buying discounted or near-expiry food."
    )
    add_table(
        doc,
        ["Mode", "Purpose"],
        [
            ["Add to my fridge", "Product at home — lookup and add to inventory"],
            ["Scan before buying", "Rescue Basket — scan discounted food, get recipe ideas"],
        ],
    )
    add_numbered(
        doc,
        [
            "Choose scan mode",
            "Use Open camera scanner or enter barcode manually",
            "Allow camera permission and align barcode inside the frame",
            "Detected barcode auto-fills and product lookup runs",
            "Review nutrition panel and safety check",
            "Fridge Scan: set expiry → Add to fridge",
            "Rescue Basket: Get recipe ideas → verdict + recommended recipes",
        ],
    )
    doc.add_paragraph("Demo barcode: 6111246721261")
    add_bullets(
        doc,
        [
            "Web: camera scanning works when browser camera permission is allowed",
            "Mobile: camera barcode scanning is included in the Scan tab",
            "If camera access fails, type the barcode manually and tap Lookup product",
        ],
    )

    doc.add_heading("10. Ingredient substitutes (Substitute tab)", level=1)
    doc.add_paragraph(
        "Search unfamiliar ingredients. Quick chips: miso, tempeh, kimchi, cassava, etc. "
        "Results show ingredient, explanation, confidence. API: GET /ingredients/{name}/similar."
    )

    doc.add_heading("11. Edit profile later", level=1)
    doc.add_paragraph(
        "Tap Profile icon in top app bar. Save changes — recommendations reload. Snackbar confirms save."
    )

    doc.add_heading("12. Product principles", level=1)
    add_table(
        doc,
        ["Type", "Examples", "Behaviour"],
        [
            ["Hard filters (safety)", "Allergies, vegetarian/vegan/halal", "Unsafe recipes excluded"],
            ["Soft signals (ranking)", "Expiry, nutrition, mood, cuisine, context", "Re-rank safe recipes"],
        ],
    )

    doc.add_heading("13. Full demo script (presentation)", level=1)
    add_numbered(
        doc,
        [
            "Open app — green API dot; note Loading ThriftyChef splash",
            "Onboarding: vegetarian, milk allergy, low sugar, Asian + Sri Lankan",
            "Fridge — add eggs, parsley, cheese, tomato; show photos and expiry bars",
            "Scan — use Open camera scanner or barcode 6111246721261 in Rescue Basket mode → recipe ideas",
            "Recipes — hero card, mood chips, food-photo cards, match % rings",
            "Toggle expiry OFF — show ranking change",
            "Mood Comfort → Healthy",
            "Open recipe — hero image, steps, why recommended",
            "Substitute — search miso",
            "Theme toggle — light/dark mode",
            "Profile icon — edit and save",
        ],
    )

    doc.add_heading("14. Demo checklist", level=1)
    for item in [
        "API health OK — green dot (service: thriftychef-api)",
        "Web app loads at 127.0.0.1:8080 (release build)",
        "ThriftyChef logo with chef hat; teal UI",
        "Onboarding centred card with sections A–E",
        "Fridge: stat cards, ingredient photos, expiry bars",
        "Scan tab: both modes, camera scan or manual entry, rescue recommendations",
        "AI recommendations: food-photo cards, match % rings",
        "Recipe detail: hero image, why-recommended card",
        "Light/dark theme toggle works",
        "Substitute search works for miso",
        "Profile editable via top bar icon",
    ]:
        p = doc.add_paragraph()
        p.add_run("☐ ").font.name = "Segoe UI Symbol"
        p.add_run(item)

    doc.add_heading("15. Troubleshooting", level=1)
    add_table(
        doc,
        ["Problem", "Fix"],
        [
            ["Blank white screen", "Use release build + Python static server (see §1)"],
            ["API not reachable", "Run python scripts/run_api.py; click Retry"],
            ["localhost:8080 fails", "Use http://127.0.0.1:8080 instead"],
            ["Camera scanner not opening", "Allow camera permission; on web use Chrome/Edge with 127.0.0.1 or HTTPS"],
            ["Camera opens but does not detect", "Improve lighting, hold barcode inside the frame, or type it manually"],
            ["No food photos", "Requires internet (Unsplash); placeholders if offline"],
            ["Old UI after update", "Rebuild web; hard refresh (Ctrl+Shift+R)"],
            ["Empty recommendations", "Wait 30s for models; add fridge items"],
            ["No safe recipes", "Relax nutrition filters or add ingredients"],
            ["Port in use", "Stop old processes on 8000/8080 and restart"],
        ],
    )

    doc.add_heading("16. Known limitations", level=1)
    add_bullets(
        doc,
        [
            "Demo user ID 5060 only — no login yet",
            "Browser camera access depends on permission settings and secure-context support",
            "Food photos from Unsplash — need internet",
            "Fridge may reset on API restart (in-memory store)",
            "Profile cached locally via SharedPreferences",
        ],
    )

    doc.add_heading("17. Screenshots for report (Appendix E)", level=1)
    add_numbered(
        doc,
        [
            "Boot / loading — Loading ThriftyChef splash",
            "Onboarding — profile card, chef-hat logo, sections A–E",
            "App shell — API dot, nav rail, theme toggle",
            "AI recommendations — hero card, food-photo cards, match rings",
            "Recipe detail — hero image, why recommended, steps",
            "Fridge — stat cards, ingredient photos, expiry bars",
            "Scan — camera scanner open, or rescue basket mode with nutrition panel",
            "Substitute — miso quick chip and result card",
            "Dark mode — Recipes tab with navy background",
        ],
    )
    doc.add_paragraph("Visual overview: deploy/ThriftyChef_Screen_Set.png")

    doc.add_heading("18. UI reference", level=1)
    add_table(
        doc,
        ["Colour", "Hex", "Use"],
        [
            ["Background (light)", "#F1F5F9", "Page background"],
            ["Primary teal", "#34A0A4", "Buttons, headings, selected chips"],
            ["Teal deep", "#268387", "Hero gradients"],
            ["Ice light", "#E6F4F5", "Hero cards, highlights"],
            ["Rose accent", "#E5A99E", "Chef wordmark and chef hat"],
            ["Warning orange", "#D97706", "Expiring items, missing chips"],
            ["Danger red", "#DC2626", "Urgent expiry, allergens"],
            ["Background (dark)", "#0B1220", "Dark mode page background"],
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    saved: list[Path] = []
    for path in (OUTPUT_V5, OUTPUT):
        try:
            doc.save(path)
            saved.append(path)
            print(f"Saved: {path}")
        except PermissionError:
            print(f"Skipped (file open?): {path}")
    if not saved:
        raise PermissionError("Could not save user guide — close open Word files and retry.")


if __name__ == "__main__":
    build()
